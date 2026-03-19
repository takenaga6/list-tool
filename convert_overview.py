from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles['Normal']
style.font.name = 'メイリオ'
style.font.size = Pt(10.5)

def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    if p.runs:
        p.runs[0].font.name = 'メイリオ'
    return p

h(doc, '企業リストアップツール 全体設計書', 0)

p = doc.add_paragraph()
r = p.add_run('サービス'); r.bold = True
p.add_run(': Offi-Stretch（理学療法士によるオフィス出張施術）')
p = doc.add_paragraph()
r = p.add_run('目的'); r.bold = True
p.add_run(': B2B営業のアタックリストを自動生成し、HubSpot CRMへ登録する')

h(doc, '全体フロー', 1)
p = doc.add_paragraph()
p.add_run(
    'python main.py を実行\n'
    '  1. 登録方式選択（確認後に登録 / 自動登録）\n'
    '  2. 検索モード選択（自動モード / 手動モード）\n'
    '  3. 期間選択（1週間〜1年）\n'
    '  4. 検索・スクレイプ・ランク判定（並列処理）\n'
    '     search_agent → scraper_agent → rank_agent\n'
    '  5. [確認モード] 候補一覧表示 → ユーザー承認\n'
    '     y=全件登録 / n=スキップ / 1,3=個別選択 / x2=除外リストへ\n'
    '  6. HubSpot登録\n'
    '  7. 完了レポート\n\n'
    '監査: python main.py --audit [--scrape]\n'
    '  → HubSpot登録済み企業を事後チェック → NG企業にフラグ + 学習'
).font.name = 'MS Gothic'

h(doc, 'ファイル構成', 1)
files = [
    ('main.py', 'エントリーポイント・UI・処理フロー制御'),
    ('config.py', '設定・定数・学習データ管理'),
    ('agents/search_agent.py', 'DuckDuckGo検索・除外フィルタ'),
    ('agents/scraper_agent.py', '企業HP情報取得・バリデーション'),
    ('agents/rank_agent.py', 'A/B/C/NG ランク判定'),
    ('agents/hubspot_agent.py', 'HubSpot API（登録・重複チェック）'),
    ('agents/hubspot_auditor.py', '登録済み企業の事後バリデーション'),
    ('agents/list_page_agent.py', 'リストページ（PDF/Excel/HTML）から一括抽出'),
    ('agents/keyword_agent.py', '検索クエリの自動生成・学習'),
    ('output/results.csv', '登録成功した企業一覧'),
    ('output/ng_list.csv', 'NG企業リスト'),
    ('output/exclude_list.csv', '除外ドメインリスト'),
    ('output/learned_exclude.json', '自動学習した除外ドメイン'),
    ('output/domain_fail_stats.json', 'ドメイン失敗カウンター（3回→自動除外）'),
    ('output/feedback.csv', 'テレアポ結果フィードバック'),
    ('output/tool.log', '実行ログ'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
hdr = t.rows[0].cells
hdr[0].text = 'ファイル'
hdr[1].text = '役割'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for f, r in files:
    row = t.add_row()
    row.cells[0].text = f
    row.cells[1].text = r
doc.add_paragraph()

h(doc, '各エージェントの役割', 1)

agents_info = [
    ('search_agent.py', [
        'DuckDuckGo で検索クエリを実行、除外フィルタを多段階で適用',
        '除外順: EXCLUDE_DOMAINS → NG_DOMAIN_KEYWORDS → 自動学習除外リスト → 海外TLD → 媒体ドメイン',
    ]),
    ('scraper_agent.py', [
        '企業HPをスクレイピングして情報抽出（会社名・代表者・住所・TEL等）',
        'check_company_fields() で5項目チェック（3/5以上で企業HP確認）',
        '失敗ドメインは record_domain_fail() でカウント → 3回で自動除外',
        'tel: リンクから電話番号を優先取得',
        '全47都道府県リストで都道府県＋市区町村を正確に抽出',
    ]),
    ('rank_agent.py', [
        '企業をA/B/C/NGにランク判定（採点方式）',
        'A=6点以上 / B=4〜5点 / C=2〜3点',
        'NG条件: 上場企業・従業員1000名超・NG業種・個人事業主等',
    ]),
    ('hubspot_agent.py', [
        'HubSpot CRM API v3 で企業を登録',
        '電話番号の国際化変換（03-1234-5678 → +81312345678）',
        '重複チェック（ドメイン・会社名）',
    ]),
    ('hubspot_auditor.py', [
        '登録済み企業を事後検証し、NGにフラグ立て（python main.py --audit）',
        'チェック順: ①ドメインキーワード → ②会社名vsドメイン → ③軽量スクレイプ（--scrape時）',
        'NG判定 → HubSpotに「【要確認】」追記 + exclude_list.csv に学習',
    ]),
    ('list_page_agent.py', [
        '健康経営優良法人リスト等の一括リストから企業を抽出',
        '対応形式: HTML / PDF / Word(.docx) / Excel(.xlsx/.xls)',
        '大規模法人ファイル（daikibo/大規模）は除外（中小企業のみ対象）',
    ]),
    ('keyword_agent.py', [
        '検索クエリのヒット率・ランク率（A率）を学習・記録',
        '有効パターンをA率順にランキングして次回実行時に提案',
    ]),
]

for name, points in agents_info:
    p = doc.add_paragraph()
    r = p.add_run(name); r.bold = True
    for pt in points:
        doc.add_paragraph('• ' + pt)

h(doc, 'ランク判定採点表', 1)
rt = doc.add_table(rows=1, cols=2)
rt.style = 'Table Grid'
hdr = rt.rows[0].cells
hdr[0].text = 'シグナル'
hdr[1].text = '点数'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for s, pt in [
    ('PR媒体掲載（KENJA GLOBAL等）', '+1'),
    ('健康経営・ウェルビーイング注力', '+1'),
    ('法定外福利厚生あり', '+1'),
    ('HPリニューアル・ブランド投資', '+1'),
    ('自社ビル・固定オフィス', '+1'),
    ('従業員数 10〜100名', '+1'),
    ('PR媒体クエリ経由', '+2 ボーナス'),
]:
    row = rt.add_row()
    row.cells[0].text = s
    row.cells[1].text = pt
doc.add_paragraph()

h(doc, '精度向上サイクル', 1)
doc.add_paragraph(
    '1. リストアップ実行\n'
    '2. 承認ステップで不要企業を除外（x[番号] → exclude_list.csv へ追加）\n'
    '3. テレアポ実施 → feedback.csv に結果記録\n'
    '4. 傾向分析 → ランク判定基準・検索パターンの調整\n'
    '5. 繰り返し（精度向上）\n\n'
    '【自動学習ループ】\n'
    '  check_company_fields 失敗 → record_domain_fail() → 3回で learned_exclude.json へ\n'
    '  承認時除外 → exclude_list.csv へ\n'
    '  監査NG → exclude_list.csv へ\n'
    '  → 次回リストアップから自動除外'
)

h(doc, 'コマンド一覧', 1)
p = doc.add_paragraph()
p.add_run(
    'cd C:\\Users\\user\\list_tool\n\n'
    '# 通常実行\n'
    'python main.py\n\n'
    '# HubSpot監査（ドメイン＋会社名チェック）\n'
    'python main.py --audit\n\n'
    '# HubSpot監査（+ 軽量スクレイプ検証）\n'
    'python main.py --audit --scrape'
).font.name = 'MS Gothic'

h(doc, 'HubSpot フィールドマッピング', 1)
ht = doc.add_table(rows=1, cols=2)
ht.style = 'Table Grid'
hdr = ht.rows[0].cells
hdr[0].text = 'HubSpotフィールド'
hdr[1].text = '取得元'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
for f, s in [
    ('会社名', 'スクレイピング'),
    ('ウェブサイト', '企業HP URL'),
    ('電話番号', 'スクレイピング → 国際化変換（+81形式）'),
    ('住所・都道府県・郵便番号', 'スクレイピング'),
    ('業種', 'スクレイピング'),
    ('従業員数', 'スクレイピング'),
    ('説明（備考）', '媒体記事URL / 検索クエリ / ランク理由'),
]:
    row = ht.add_row()
    row.cells[0].text = f
    row.cells[1].text = s

out_path = r'C:\Users\user\list_tool\OVERVIEW.docx'
doc.save(out_path)
print('保存完了: ' + out_path)
